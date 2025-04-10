import re
import shutil
from pathlib import Path
import docx # type: ignore
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
import copy
# --- ДОБАВЛЕН ИМПОРТ ---
from docx.enum.text import WD_ALIGN_PARAGRAPH
# -----------------------

class DocxHandler:
    """
    Класс для инкапсуляции операций с файлами DOCX.
    (Версия 4.1: Исправлен NameError)
    """
    KEY_PATTERN = re.compile(r"\{\{.*?\}\}")
    DYNAMIC_TABLE_PATTERN = re.compile(r"\{\{DYNAMIC_TABLE::(\w+)\}\}")

    # --- Метод find_keys_in_template остается как в v2 ---
    def _get_paragraph_keys(self, paragraph) -> set[str]:
        full_para_text = "".join(run.text for run in paragraph.runs)
        return set(self.KEY_PATTERN.findall(full_para_text))

    def find_keys_in_template(self, docx_path: Path) -> dict:
        found_keys = set()
        found_tables_info: dict[str, dict] = {}
        try:
            document = docx.Document(docx_path)
            for para in document.paragraphs:
                found_keys.update(self._get_paragraph_keys(para))
            for table in document.tables:
                for row_idx, row in enumerate(table.rows):
                    table_id_found_in_row: str | None = None
                    # Ищем маркер только в первой ячейке для определения строки-шаблона
                    if row.cells: # Проверка, что ячейки существуют
                        first_cell_text = "".join(p.text for p in row.cells[0].paragraphs)
                        match = self.DYNAMIC_TABLE_PATTERN.search(first_cell_text)
                        if match:
                            table_id_found_in_row = match.group(1)
                            # Собираем ключи из остальных ячеек этой строки
                            if table_id_found_in_row not in found_tables_info:
                                ordered_template_keys = []
                                for cell_idx in range(1, len(row.cells)): # Начиная со второй ячейки
                                    cell = row.cells[cell_idx]
                                    for para in cell.paragraphs:
                                        keys_in_para = self._get_paragraph_keys(para)
                                        for key in keys_in_para:
                                            # Не добавляем сам маркер и избегаем дубликатов
                                            if not self.DYNAMIC_TABLE_PATTERN.match(key) and key not in ordered_template_keys:
                                                ordered_template_keys.append(key)
                                found_tables_info[table_id_found_in_row] = {'template_keys': ordered_template_keys}
                                print(f"Найдена таблица '{table_id_found_in_row}' с template_keys: {ordered_template_keys}")

                    # Собираем ВСЕ ключи из таблицы в общий список
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            found_keys.update(self._get_paragraph_keys(para))


            table_markers_full = {f"{{{{DYNAMIC_TABLE::{tid}}}}}" for tid in found_tables_info}
            keys_only = found_keys - table_markers_full
        except Exception as e:
            print(f"Ошибка при чтении или обработке файла {docx_path}: {e}")
            return {'keys': set(), 'tables': {}}
        return {'keys': keys_only, 'tables': found_tables_info}
    # --- Конец метода find_keys_in_template ---

    # --- Метод _replace_text_in_paragraph остается как в v3 ---
    def _replace_text_in_paragraph(self, paragraph: Paragraph, old_text: str, new_text: str):
        if old_text not in paragraph.text: return
        runs_info = [(run, run.text) for run in paragraph.runs]
        full_text = "".join(info[1] for info in runs_info)
        start_index = full_text.find(old_text)
        if start_index == -1: return
        end_index = start_index + len(old_text)
        current_pos = 0; first_run_idx = -1; last_run_idx = -1; style_run = None
        for i, (run, text) in enumerate(runs_info):
            run_len = len(text); run_start = current_pos; run_end = current_pos + run_len
            if run_start < end_index and run_end > start_index:
                if first_run_idx == -1: first_run_idx = i; style_run = run
                last_run_idx = i
                replace_start_in_run = max(0, start_index - run_start)
                replace_end_in_run = min(run_len, end_index - run_start)
                original_run_text = text; new_run_text = ""
                if i == first_run_idx: new_run_text = original_run_text[:replace_start_in_run] + new_text
                elif i > first_run_idx and i < last_run_idx: new_run_text = ""
                if i == last_run_idx: new_run_text += original_run_text[replace_end_in_run:]
                run.text = new_run_text
            current_pos += run_len
        if first_run_idx != -1 and last_run_idx != -1:
            for i in range(last_run_idx, first_run_idx, -1):
                 run_to_check = paragraph.runs[i]
                 if not run_to_check.text:
                     p = run_to_check._element.getparent()
                     if p is not None: p.remove(run_to_check._element)
    # --- Конец метода _replace_text_in_paragraph ---

    # --- Метод _copy_cell_formatting остается как в v4 ---
    def _copy_cell_formatting(self, source_cell, target_cell):
        if source_cell._tc.tcPr:
            target_tcPr = target_cell._tc.get_or_add_tcPr()
            if source_cell._tc.tcPr.vAlign:
                 vAlign = OxmlElement('w:vAlign'); vAlign.val = source_cell._tc.tcPr.vAlign.val
                 target_tcPr.append(vAlign)
        if source_cell.paragraphs and target_cell.paragraphs:
            source_p = source_cell.paragraphs[0]; target_p = target_cell.paragraphs[0]
            target_p.alignment = source_p.alignment; target_p.style = source_p.style
            if source_p.runs and target_p.runs:
                source_r = source_p.runs[0]; target_r = target_p.runs[0]
                target_r.bold = source_r.bold; target_r.italic = source_r.italic
                target_r.underline = source_r.underline; target_r.font.name = source_r.font.name
                target_r.font.size = source_r.font.size; target_r.font.color.rgb = source_r.font.color.rgb
    # --- Конец метода _copy_cell_formatting ---

    # --- Метод generate_document остается как в v4 (но теперь импорт есть) ---
    def generate_document(self, template_path: Path, output_path: Path, project_keys_data: dict) -> bool:
        print(f"Генерация документа из '{template_path.name}' в '{output_path}'...")
        try:
            shutil.copy2(template_path, output_path)
            print(f"Шаблон скопирован в {output_path}")
            doc = docx.Document(output_path)
            key_value_map = {}; table_definitions = {}
            for key_id, data in project_keys_data.items():
                if data.get('type') == 'dynamic_table': table_definitions[key_id] = data
                else: key_value_map[key_id] = data.get('value', '')
            print("Предварительная замена простых ключей в параграфах...")
            for para in doc.paragraphs:
                inline_keys = self.KEY_PATTERN.findall(para.text)
                for key in inline_keys:
                    if key in key_value_map: self._replace_text_in_paragraph(para, key, key_value_map[key])
            print(f"Обработка {len(table_definitions)} динамических таблиц...")
            processed_table_elements = set()
            for table_idx, table in enumerate(doc.tables):
                if table._tbl in processed_table_elements: continue
                template_row_index = -1; table_id_to_process = None; table_template_keys = None
                for r_idx, row in enumerate(table.rows):
                    if row.cells:
                        first_cell_text = "".join(p.text for p in row.cells[0].paragraphs)
                        match = self.DYNAMIC_TABLE_PATTERN.search(first_cell_text)
                        if match:
                            potential_table_id = match.group(1)
                            if potential_table_id in table_definitions:
                                table_id_to_process = potential_table_id; template_row_index = r_idx
                                table_template_keys = table_definitions[table_id_to_process].get('template_keys', [])
                                print(f"Найдена таблица '{table_id_to_process}' в документе (индекс {table_idx}, строка-шаблон {template_row_index}).")
                                break
                if table_id_to_process and template_row_index != -1 and table_template_keys is not None:
                    processed_table_elements.add(table._tbl)
                    table_data_rows = table_definitions[table_id_to_process].get('data', [])
                    template_row = table.rows[template_row_index]
                    print(f"Очистка строк данных в таблице '{table_id_to_process}'...")
                    for r_idx in range(len(table.rows) - 1, 0, -1):
                        row_to_remove = table.rows[r_idx]; table._tbl.remove(row_to_remove._tr)
                    print(f"Добавление {len(table_data_rows)} строк в таблицу '{table_id_to_process}'...")
                    num_columns = len(table.rows[0].cells) if table.rows else 0
                    for data_idx, row_data_dict in enumerate(table_data_rows):
                        new_row_obj = table.add_row(); new_row_cells = new_row_obj.cells
                        if len(new_row_cells) != num_columns: print(f"Предупреждение: Несовпадение кол-ва ячеек..."); continue
                        cell_num = new_row_cells[0]; cell_num.text = str(data_idx + 1)
                        if len(template_row.cells) > 0:
                             self._copy_cell_formatting(template_row.cells[0], cell_num)
                             # Используем импортированный WD_ALIGN_PARAGRAPH
                             if cell_num.paragraphs: cell_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for key_idx, template_key in enumerate(table_template_keys):
                            cell_idx_in_doc = key_idx + 1
                            if cell_idx_in_doc < len(new_row_cells):
                                cell_data = new_row_cells[cell_idx_in_doc]
                                value = str(row_data_dict.get(template_key, ''))
                                cell_data.text = value
                                if cell_idx_in_doc < len(template_row.cells):
                                    self._copy_cell_formatting(template_row.cells[cell_idx_in_doc], cell_data)
            print("Финальная замена простых ключей...")
            all_paragraphs = list(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells: all_paragraphs.extend(cell.paragraphs)
            for para in all_paragraphs:
                keys_to_replace = list(key_value_map.keys())
                for key in keys_to_replace:
                     if key in para.text: self._replace_text_in_paragraph(para, key, key_value_map[key])
            doc.save(output_path)
            print(f"Документ успешно сгенерирован и сохранен: {output_path}")
            return True
        except FileNotFoundError: print(f"Ошибка: Шаблон не найден {template_path}"); return False
        except Exception as e:
            print(f"Ошибка при генерации документа {output_path}: {e}")
            if output_path.exists():
                try: output_path.unlink()
                except OSError: pass
            return False

# --- Блок if __name__ == '__main__' остается для тестов ---
if __name__ == '__main__':
    # Импорт WD_ALIGN_PARAGRAPH здесь уже был для теста, но он нужен и выше
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    handler = DocxHandler()
    tpl_path = Path("E:/docx_dormatter/ТП++.docx") # Пример пути
    proj_path = Path("E:/docx_dormatter/Тестпроект.dfp") # Пример пути
    out_path = Path("E:/docx_dormatter/generated_docs/ТП++_generated_TABLE.docx")
    if not tpl_path.exists(): print(f"Файл шаблона не найден: {tpl_path}")
    elif not proj_path.exists(): print(f"Файл проекта не найден: {proj_path}")
    else:
        import json
        try:
            with open(proj_path, 'r', encoding='utf-8') as f: project_data = json.load(f)
            keys_data = project_data.get('keys_data', {})
            out_path.parent.mkdir(parents=True, exist_ok=True)
            handler.generate_document(tpl_path, out_path, keys_data)
        except Exception as e: print(f"Ошибка при тестовой генерации: {e}")